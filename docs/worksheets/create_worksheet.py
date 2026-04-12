"""Generate the Human-AI Collaboration Worksheet for Mirofish development.

Run: python docs/worksheets/create_worksheet.py
Output: docs/worksheets/mirofish_human_worksheet.docx
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def add_heading(doc, text, level=1):
    heading = doc.add_heading(text, level=level)
    return heading


def add_checkbox_item(doc, text, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(indent * 0.5)
    run = p.add_run("[ ] " + text)
    run.font.size = Pt(11)
    return p


def add_fillable_field(doc, label, lines=3):
    p = doc.add_paragraph()
    run = p.add_run(label + ":")
    run.bold = True
    run.font.size = Pt(11)
    for _ in range(lines):
        p = doc.add_paragraph()
        run = p.add_run("_" * 80)
        run.font.color.rgb = RGBColor(200, 200, 200)


def add_table_row(table, cells_text):
    row = table.add_row()
    for i, text in enumerate(cells_text):
        row.cells[i].text = text


def create_worksheet():
    doc = Document()

    # Title
    title = doc.add_heading("Mirofish Prediction Engine", level=0)
    subtitle = doc.add_heading("Human-AI Collaboration Worksheet", level=1)
    doc.add_paragraph(
        "This workbook is designed for structured human-AI collaboration in developing "
        "the Mirofish prediction engine. Complete each section to provide the AI system "
        "with your domain knowledge, preferences, and strategic decisions."
    )
    doc.add_paragraph(
        "Date: ____________  |  Version: 1.0  |  Human Lead: _____________________"
    )
    doc.add_paragraph("")

    # =========================================================================
    # SECTION 1: STRATEGIC VISION
    # =========================================================================
    add_heading(doc, "Section 1: Strategic Vision & Goals", level=1)

    doc.add_paragraph(
        "Define what you want to achieve with Mirofish. "
        "Your answers shape the system's priorities."
    )

    add_fillable_field(doc, "1.1 What is your primary goal? (circle one or write your own)")
    doc.add_paragraph(
        "   a) Learn about quant trading and prediction markets\n"
        "   b) Build a working prototype for personal use\n"
        "   c) Generate consistent returns on Polymarket\n"
        "   d) Create a system to share with coworkers\n"
        "   e) Other: _______________"
    )

    add_fillable_field(doc, "1.2 What prediction market platforms do you want to target?")
    add_fillable_field(doc, "1.3 What is your initial bankroll (paper trading budget)?")
    add_fillable_field(doc, "1.4 What is your risk tolerance? (conservative / moderate / aggressive)")
    add_fillable_field(doc, "1.5 What market categories interest you most? (politics, crypto, sports, economics, tech, etc.)")
    add_fillable_field(doc, "1.6 What is your time horizon? (daily trading / weekly / monthly / event-based)")

    # =========================================================================
    # SECTION 2: DOMAIN KNOWLEDGE INJECTION
    # =========================================================================
    add_heading(doc, "Section 2: Domain Knowledge Injection", level=1)

    doc.add_paragraph(
        "The AI system needs YOUR domain expertise to make better predictions. "
        "What do you know that most people don't?"
    )

    add_fillable_field(doc, "2.1 What domains do you have deep expertise in?", 4)
    add_fillable_field(doc, "2.2 What news sources do you trust most? (specific outlets, analysts, feeds)", 4)
    add_fillable_field(doc, "2.3 What biases do you know you have? (political leaning, tech optimism, etc.)", 3)
    add_fillable_field(doc, "2.4 What 'hidden signals' do you watch that most people ignore?", 4)
    add_fillable_field(doc, "2.5 What predictions have you made recently? How did they turn out?", 5)

    # =========================================================================
    # SECTION 3: QUANT FINANCE FOUNDATIONS
    # =========================================================================
    add_heading(doc, "Section 3: Quant Finance Knowledge Assessment", level=1)

    doc.add_paragraph(
        "Rate your familiarity with each concept (1=none, 5=expert). "
        "This helps the AI calibrate explanations and suggestions."
    )

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Concept"
    hdr[1].text = "Familiarity (1-5)"
    hdr[2].text = "Notes / Questions"

    concepts = [
        "Kelly Criterion (position sizing)",
        "Bayesian probability updating",
        "Mean reversion vs. trend following",
        "Market microstructure (order books, spreads)",
        "Time series analysis (ARIMA, GARCH)",
        "Signal processing for finance",
        "Game theory (zero-sum, Nash equilibrium)",
        "Portfolio optimization (Markowitz)",
        "Derivatives pricing (Black-Scholes)",
        "Risk management (VaR, CVaR, drawdown)",
        "Information theory (entropy, KL divergence)",
        "Calibration (Brier score, ECE)",
        "Prediction market theory (Hayek, info aggregation)",
        "Algorithmic trading strategies",
        "Network/graph analysis",
        "Machine learning for finance",
        "LLM/NLP for sentiment analysis",
        "Blockchain / DeFi mechanics",
    ]

    for concept in concepts:
        add_table_row(table, [concept, "", ""])

    doc.add_paragraph("")
    add_fillable_field(doc, "3.1 Which classic textbooks or papers have you read?", 4)
    add_fillable_field(doc, "3.2 What concepts do you want to learn more about?", 3)

    # =========================================================================
    # SECTION 4: MARKET SELECTION CRITERIA
    # =========================================================================
    add_heading(doc, "Section 4: Market Selection Criteria", level=1)

    doc.add_paragraph(
        "Help the system choose which markets to analyze."
    )

    add_fillable_field(doc, "4.1 Minimum market volume threshold (USD)")
    add_fillable_field(doc, "4.2 Minimum market liquidity threshold (USD)")
    add_fillable_field(doc, "4.3 Categories to INCLUDE (e.g., politics, crypto, science)")
    add_fillable_field(doc, "4.4 Categories to EXCLUDE (e.g., sports, entertainment)")
    add_fillable_field(doc, "4.5 Time-to-resolution preference (days to weeks / weeks to months / months to years)")
    add_fillable_field(doc, "4.6 Minimum edge threshold before betting (%)")

    # =========================================================================
    # SECTION 5: RISK TOLERANCE WORKSHEET
    # =========================================================================
    add_heading(doc, "Section 5: Risk Management Parameters", level=1)

    doc.add_paragraph(
        "Define your risk boundaries. The AI will NEVER exceed these."
    )

    table2 = doc.add_table(rows=1, cols=3)
    table2.style = "Table Grid"
    hdr2 = table2.rows[0].cells
    hdr2[0].text = "Parameter"
    hdr2[1].text = "Your Value"
    hdr2[2].text = "Default"

    params = [
        ("Total bankroll (USD)", "", "$1,000"),
        ("Max single position (% of bankroll)", "", "5%"),
        ("Max drawdown before stopping (%)", "", "15%"),
        ("Kelly fraction (1.0=full, 0.25=quarter)", "", "0.25"),
        ("Min edge to trade (%)", "", "5%"),
        ("Paper trading only? (yes/no)", "", "yes"),
        ("Max concurrent positions", "", "10"),
        ("Daily loss limit (USD)", "", "No limit"),
    ]

    for param in params:
        add_table_row(table2, list(param))

    # =========================================================================
    # SECTION 6: SWARM CONFIGURATION
    # =========================================================================
    add_heading(doc, "Section 6: Swarm Intelligence Configuration", level=1)

    doc.add_paragraph(
        "Configure which Fish agents to activate and their priorities."
    )

    table3 = doc.add_table(rows=1, cols=4)
    table3.style = "Table Grid"
    hdr3 = table3.rows[0].cells
    hdr3[0].text = "Fish Agent"
    hdr3[1].text = "Activate? (Y/N)"
    hdr3[2].text = "Priority (1-10)"
    hdr3[3].text = "Notes"

    fish = [
        ("Geopolitical Analyst", "", "", "International relations, political risk"),
        ("Financial Quant", "", "", "Statistics, pricing, base rates"),
        ("Bayesian Statistician", "", "", "Prior updating, uncertainty"),
        ("Investigative Journalist", "", "", "Hidden info, primary sources"),
        ("Contrarian Thinker", "", "", "Devil's advocate, bias detection"),
        ("Domain Expert", "", "", "Specialized subject knowledge"),
        ("Calibration Specialist", "", "", "Meta-analysis, bias correction"),
        ("(Custom 1): ________", "", "", ""),
        ("(Custom 2): ________", "", "", ""),
    ]

    for f in fish:
        add_table_row(table3, list(f))

    # =========================================================================
    # SECTION 7: HUMAN-AI INTERACTION PREFERENCES
    # =========================================================================
    add_heading(doc, "Section 7: Human-AI Interaction Preferences", level=1)

    add_fillable_field(doc, "7.1 How often do you want the AI to ask for your input? (every decision / major decisions only / minimal)")
    add_fillable_field(doc, "7.2 How detailed should AI explanations be? (brief / moderate / comprehensive)")
    add_fillable_field(doc, "7.3 What format do you prefer for reports? (bullet points / narrative / tables / visual)")
    add_fillable_field(doc, "7.4 How should the AI handle disagreements with your input?")
    add_fillable_field(doc, "7.5 What languages do you want the AI to use? (English / Korean / both)")

    # =========================================================================
    # SECTION 8: LEARNING GOALS
    # =========================================================================
    add_heading(doc, "Section 8: Learning & Development Goals", level=1)

    doc.add_paragraph(
        "What do you want to LEARN through building this system?"
    )

    add_checkbox_item(doc, "Quantitative finance fundamentals (Bachelier, Markowitz, Thorp)")
    add_checkbox_item(doc, "Kelly Criterion and optimal betting theory")
    add_checkbox_item(doc, "Time series analysis and signal processing")
    add_checkbox_item(doc, "Game theory applied to markets")
    add_checkbox_item(doc, "LLM/AI applied to financial prediction")
    add_checkbox_item(doc, "Prediction market theory and information aggregation")
    add_checkbox_item(doc, "Python programming for quantitative analysis")
    add_checkbox_item(doc, "Network science and graph analysis")
    add_checkbox_item(doc, "Risk management and portfolio theory")
    add_checkbox_item(doc, "Algorithmic trading strategies")
    add_checkbox_item(doc, "Blockchain and DeFi mechanics")
    add_checkbox_item(doc, "Other: _________________________________")

    doc.add_paragraph("")
    add_fillable_field(doc, "8.1 What journal articles should be prioritized for the literature review?", 4)

    doc.add_paragraph(
        "Recommended journals (from Wikipedia quant finance article):\n"
        "- SIAM Journal on Financial Mathematics\n"
        "- Journal of Portfolio Management\n"
        "- Quantitative Finance\n"
        "- Risk Magazine\n"
        "- Wilmott Magazine\n"
        "- Finance and Stochastics\n"
        "- Mathematical Finance"
    )

    # =========================================================================
    # SECTION 9: READING LIST HOMEWORK
    # =========================================================================
    add_heading(doc, "Section 9: Reading List (Human Homework)", level=1)

    doc.add_paragraph(
        "These readings will deepen your understanding and help you provide "
        "better domain knowledge to the AI system. Check off as you complete them."
    )

    add_heading(doc, "Foundational Classics", level=2)
    add_checkbox_item(doc, "Kelly (1956) 'A New Interpretation of Information Rate' -- THE betting theory paper")
    add_checkbox_item(doc, "Thorp (1962) 'Beat the Dealer' -- Kelly criterion in practice")
    add_checkbox_item(doc, "Markowitz (1952) 'Portfolio Selection' -- Modern portfolio theory")
    add_checkbox_item(doc, "Hayek (1945) 'The Use of Knowledge in Society' -- Information aggregation")
    add_checkbox_item(doc, "Wolfers & Zitzewitz (2004) 'Prediction Markets' -- Prediction market theory")

    add_heading(doc, "Modern LLM + Finance (our direct foundation)", level=2)
    add_checkbox_item(doc, "Schoenegger et al. (2024) 'Wisdom of the Silicon Crowd' -- LLM ensemble = human crowd")
    add_checkbox_item(doc, "Halawi et al. (2024) 'Approaching Human-Level Forecasting' -- RAG pipeline for prediction")
    add_checkbox_item(doc, "Barot & Borkhatariya (2026) 'PolySwarm' -- Our primary architecture reference")
    add_checkbox_item(doc, "Turtel et al. (2025) 'Outcome-based RL' -- RL fine-tuning for Polymarket")
    add_checkbox_item(doc, "Saguillo et al. (2025) 'Empirical Arbitrage on Polymarket' -- $40M arbitrage analysis")
    add_checkbox_item(doc, "Tsang & Yang (2026) 'Anatomy of Polymarket' -- Market microstructure")

    add_heading(doc, "Signal Processing & Time Series", level=2)
    add_checkbox_item(doc, "Hamilton (1994) 'Time Series Analysis' -- THE time series textbook")
    add_checkbox_item(doc, "Cont (2001) 'Empirical Properties of Asset Returns' -- Stylized facts")
    add_checkbox_item(doc, "Lopez-Lira & Tang (2025) 'ChatGPT Stock Price Movements' -- LLM sentiment trading")

    add_heading(doc, "Game Theory & Algorithmic Trading", level=2)
    add_checkbox_item(doc, "Easley & Kleinberg (2010) 'Networks, Crowds, and Markets' -- FREE online textbook")
    add_checkbox_item(doc, "De Prado (2018) 'Advances in Financial Machine Learning' -- ML for quant trading")
    add_checkbox_item(doc, "Chan (2009) 'Quantitative Trading' -- Practical algorithmic trading")

    # =========================================================================
    # SECTION 10: FEEDBACK & EVOLUTION
    # =========================================================================
    add_heading(doc, "Section 10: Feedback & System Evolution", level=1)

    doc.add_paragraph(
        "After each development session, record your feedback here. "
        "This drives the system's evolution."
    )

    for i in range(1, 6):
        add_heading(doc, f"Session {i} Feedback", level=2)
        doc.add_paragraph(f"Date: ____________")
        add_fillable_field(doc, "What worked well?", 2)
        add_fillable_field(doc, "What should change?", 2)
        add_fillable_field(doc, "New ideas or features to add", 2)
        add_fillable_field(doc, "Domain knowledge to inject (what did you learn?)", 2)

    # =========================================================================
    # SECTION 11: QUANT AREAS DEEP DIVE
    # =========================================================================
    add_heading(doc, "Section 11: Quant Finance Areas Assessment", level=1)

    doc.add_paragraph(
        "From Wikipedia: Areas of work in quantitative finance. "
        "Rate relevance to Mirofish and your interest level."
    )

    table4 = doc.add_table(rows=1, cols=4)
    table4.style = "Table Grid"
    hdr4 = table4.rows[0].cells
    hdr4[0].text = "Area"
    hdr4[1].text = "Relevance (1-5)"
    hdr4[2].text = "Interest (1-5)"
    hdr4[3].text = "How it applies to Mirofish"

    areas = [
        "Trading strategy development",
        "Portfolio management & optimization",
        "Derivatives pricing & hedging",
        "Risk management",
        "Credit analysis",
        "Asset & liability management",
        "Structured finance",
        "Asset pricing",
        "Algorithmic trading",
        "Market making",
        "Statistical arbitrage",
    ]

    for area in areas:
        add_table_row(table4, [area, "", "", ""])

    # Save
    output_path = "docs/worksheets/mirofish_human_worksheet.docx"
    doc.save(output_path)
    print(f"Worksheet saved to: {output_path}")
    return output_path


if __name__ == "__main__":
    create_worksheet()
