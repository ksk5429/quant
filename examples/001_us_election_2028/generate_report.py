"""Generate Exercise 001 Report: 2028 US Election Markets."""

import json
import sys
from datetime import datetime
from pathlib import Path

sys.path.insert(0, str(Path(__file__).parent.parent.parent))

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_table_from_data(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.runs[0].bold = True
            p.runs[0].font.size = Pt(9)
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = str(val)
            for p in row.cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    return table


def main():
    doc = Document()

    # Title
    title = doc.add_heading("Mirofish Exercise 001", level=0)
    doc.add_heading("2028 US Election — Prediction Market Analysis", level=1)

    # Metadata
    meta = doc.add_paragraph()
    meta.add_run("Date: ").bold = True
    meta.add_run("2026-04-12\n")
    meta.add_run("Markets Analyzed: ").bold = True
    meta.add_run("15 (Polymarket)\n")
    meta.add_run("Fish Agents: ").bold = True
    meta.add_run("3 (Geopolitical, Quant, Contrarian)\n")
    meta.add_run("Total Analyses: ").bold = True
    meta.add_run("45 (3 Fish x 15 markets)\n")
    meta.add_run("Aggregation: ").bold = True
    meta.add_run("Bayesian confidence-weighted fusion\n")
    meta.add_run("Position Sizing: ").bold = True
    meta.add_run("Quarter-Kelly criterion ($1,000 paper bankroll)\n")
    meta.add_run("Result: ").bold = True
    meta.add_run("0 actionable trades (all PASS — edges below 5% threshold)")

    # Executive Summary
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(
        "The Mirofish swarm analyzed 15 active Polymarket markets related to the "
        "2028 US Presidential Election. All markets are ultra-low-probability tail "
        "candidates (implied probability 0.5-0.9%), including 13 Democratic nomination "
        "markets and 2 general election markets."
    )
    doc.add_paragraph(
        "Result: No actionable trades. The swarm correctly identified that edges in "
        "these markets are too small (1-2 cents absolute) to overcome transaction "
        "costs and capital lockup. This is the disciplined outcome — the system's "
        "value is in knowing when NOT to trade."
    )

    # Key Findings
    doc.add_heading("Key Findings", level=1)

    findings = [
        ("Constitutional Ineligibility — MrBeast",
         "All 3 Fish independently detected that MrBeast (born 1998) would be 30 "
         "on Election Day 2028, below the constitutional requirement of 35. The "
         "market at 0.7c should theoretically be 0c. This represents the clearest "
         "inefficiency in the dataset, but the edge is too small for practical trading."),
        ("Real Politicians Underpriced vs. Celebrities",
         "Phil Murphy (sitting NJ governor) and Tim Walz (2024 VP nominee) are "
         "priced in the same 0.6-0.75c band as Kim Kardashian and MrBeast. The "
         "swarm estimates Murphy at 1.4% and Walz at 1.6% — roughly 2x the "
         "market price. These are the most underpriced markets in the dataset."),
        ("Cross-Market Anomaly — Walz",
         "Fish Quant detected that Walz general election (0.65c) vs. nomination "
         "(0.75c) implies P(general|nomination) = 87%, which is historically "
         "too high for any non-frontrunner. The general election market appears "
         "overpriced relative to the nomination market."),
        ("Market Structure — Lazy Pricing",
         "All 15 tail candidates are priced in a narrow 0.55-0.95c band regardless "
         "of structural viability. The market does not distinguish between 'implausible "
         "but possible' (Murphy, Walz) and 'structurally impossible' (MrBeast, Sanders "
         "at age 87). This lazy pricing creates micro-inefficiencies."),
        ("Probability Space Gap",
         "The sum of all 15 YES prices is ~10.5c, meaning these candidates collectively "
         "account for ~10.5% of the probability space. The remaining ~89.5% is assigned "
         "to frontrunners not in this dataset (Newsom, Shapiro, Whitmer, etc.)."),
    ]

    for title_text, body_text in findings:
        doc.add_heading(title_text, level=2)
        doc.add_paragraph(body_text)

    # Data Table
    doc.add_heading("Swarm Analysis Results", level=1)
    doc.add_paragraph(
        "The following table shows each market's current price, individual Fish "
        "estimates, and the Bayesian-aggregated swarm consensus. All probabilities "
        "are expressed as percentages."
    )

    headers = ["Candidate", "Mkt %", "Geo %", "Quant %", "Contra %", "Swarm %", "Verdict"]
    rows = [
        ["Tim Walz (nom)", "0.75", "2.00", "0.70", "2.50", "1.60", "Underpriced"],
        ["Phil Murphy", "0.65", "1.50", "0.80", "2.00", "1.40", "Underpriced"],
        ["Tim Walz (pres)", "0.65", "1.00", "0.50", "1.20", "0.90", "Slightly under"],
        ["Andrew Yang", "0.75", "0.50", "0.60", "1.50", "0.87", "Monitor"],
        ["Beto O'Rourke", "0.65", "0.80", "0.70", "0.80", "0.77", "Fairly priced"],
        ["Oprah Winfrey", "0.95", "0.20", "0.60", "0.70", "0.50", "Overpriced"],
        ["Liz Cheney", "0.85", "0.20", "0.60", "0.40", "0.40", "Overpriced"],
        ["Chelsea Clinton", "0.85", "0.20", "0.50", "0.50", "0.40", "Overpriced"],
        ["George Clooney", "0.75", "0.10", "0.30", "0.80", "0.40", "Fairly priced"],
        ["Hillary Clinton", "0.75", "0.30", "0.50", "0.30", "0.37", "Overpriced"],
        ["LeBron (nom)", "0.75", "0.10", "0.30", "0.50", "0.30", "Overpriced"],
        ["Bernie Sanders", "0.55", "0.20", "0.30", "0.20", "0.23", "Overpriced"],
        ["MrBeast", "0.75", "0.05", "0.20", "0.30", "0.18", "Ineligible"],
        ["Kim Kardashian", "0.65", "0.05", "0.20", "0.30", "0.18", "Overpriced"],
        ["LeBron (pres)", "0.55", "0.05", "0.30", "0.20", "0.18", "Overpriced"],
    ]

    add_table_from_data(doc, headers, rows)

    # Fish Performance
    doc.add_heading("Individual Fish Performance", level=1)

    doc.add_heading("Fish Geopolitical Analyst", level=2)
    doc.add_paragraph(
        "Focused on institutional pathways, political infrastructure, and historical "
        "precedent. Produced the most conservative estimates for celebrities (0.05-0.10%) "
        "and the most differentiated estimates for real politicians (1.0-2.0% for Murphy/Walz). "
        "Strongest analysis: correctly identifying that Andrew Yang left the Democratic Party "
        "and would need to rejoin, reducing his probability."
    )

    doc.add_heading("Fish Quant", level=2)
    doc.add_paragraph(
        "Applied base rate analysis, volume/liquidity metrics, and cross-market consistency "
        "checks. Key contribution: detecting the Walz general-vs-nomination pricing anomaly "
        "(implied conditional probability of 87%). Also identified MrBeast's constitutional "
        "ineligibility through age calculation. Highest confidence scores across all analyses "
        "(0.72-0.92), reflecting quantitative certainty."
    )

    doc.add_heading("Fish Contrarian", level=2)
    doc.add_paragraph(
        "Challenged consensus on every market, looking for reasons prices might be wrong. "
        "Produced the highest estimates for Murphy (2.0%), Walz (2.5%), and Yang (1.5%) — "
        "all real politicians the market undervalues. Also correctly argued that being "
        "contrarian on MrBeast and Sanders would be wrong — sometimes the consensus IS right. "
        "Lowest confidence scores (0.50-0.65 for candidates it rated higher), appropriately "
        "reflecting uncertainty in contrarian positions."
    )

    # Critique & Improvements
    doc.add_heading("Self-Critique & Improvements for Next Exercise", level=1)

    doc.add_heading("What Worked Well", level=2)
    for item in [
        "3 Fish with distinct personas produced genuinely different analyses",
        "Cross-market anomaly detection (Walz nom vs general) — novel insight",
        "Constitutional ineligibility catch (MrBeast) — all 3 Fish found it independently",
        "Disciplined PASS on all markets — correct decision given tiny edges",
        "Full pipeline worked end-to-end: scan -> analyze -> aggregate -> signal",
        "File-based communication via shared_state/ worked flawlessly",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("What Should Improve", level=2)
    for item in [
        "All 15 markets were ultra-low-probability tail candidates — not ideal for "
        "demonstrating edge detection. Next exercise should target 20-80% probability markets.",
        "No news/RAG context provided to Fish — they analyzed from knowledge cutoff only. "
        "Adding real-time news would improve analysis quality.",
        "Only 3 Fish used. Adding Bayesian Statistician and Domain Expert would improve diversity.",
        "No debate round (Strategy C) — Fish analyzed independently without challenging "
        "each other. Adding adversarial pairing would sharpen estimates.",
        "Market prices were visible in some Fish analyses (the CLAUDE.md instruction to "
        "withhold prices was not consistently followed). Need stricter price-blinding.",
        "Aggregation used simple confidence weighting — no historical Brier scores yet "
        "(first exercise, no track record). Calibration will improve with more exercises.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    # Next Exercises
    doc.add_heading("Proposed Next Exercises", level=1)

    exercises = [
        ("Exercise 002: Crypto Markets",
         "Analyze Bitcoin, Ethereum, and crypto regulatory markets. These tend to have "
         "higher probabilities (20-80%) and larger edges. Good for testing the Kelly "
         "criterion with meaningful position sizes."),
        ("Exercise 003: Geopolitics & Conflict",
         "Russia-Ukraine ceasefire, Taiwan scenarios, trade war escalation. High-impact "
         "events where geopolitical Fish has strongest edge."),
        ("Exercise 004: Economic Indicators",
         "Fed rate decisions, inflation targets, recession probability. Quant Fish's "
         "home territory — base rates and time series are most useful here."),
        ("Exercise 005: Technology & AI",
         "AI regulation, tech company milestones, product launches. Domain Expert Fish "
         "would add the most value here."),
        ("Exercise 006: Sports (Calibration Test)",
         "Sports markets have fast resolution and clear outcomes — ideal for building "
         "calibration data quickly. 50+ resolutions in one week is possible."),
    ]

    for title_text, body_text in exercises:
        doc.add_heading(title_text, level=2)
        doc.add_paragraph(body_text)

    # Methodology Reference
    doc.add_heading("Methodology", level=1)
    doc.add_paragraph(
        "Aggregation: Bayesian confidence-weighted fusion. "
        "P_swarm = sum(w_i * P_i) / sum(w_i), where w_i = confidence_i. "
        "Position sizing: Quarter-Kelly criterion with 5% max position cap "
        "and 15% drawdown stop. Paper trading mode (no real capital at risk). "
        "See docs/MATHEMATICAL_FOUNDATIONS.md for full derivation."
    )

    # Footer
    doc.add_paragraph("")
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(
        "Generated by Mirofish Prediction Engine v0.1.0\n"
        "Human-AI Collaboration: KSK (human) + Claude Code Opus 4.6 (AI)\n"
        f"Report date: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
    )
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(128, 128, 128)

    # Save
    out = Path(__file__).parent / "exercise_001_report.docx"
    doc.save(str(out))
    print(f"Report saved: {out}")


if __name__ == "__main__":
    main()
